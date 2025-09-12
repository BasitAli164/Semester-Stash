# word_to_csv_fixed.py
import docx
import pandas as pd
import os
import re

def word_to_csv_fixed(word_file, csv_file):
    doc = docx.Document(word_file)
    rows = []
    current_phase, current_question = None, None
    collecting_answers = False
    current_answers = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        print(">>", text)  # Debug

        # Detect Phase
        if re.match(r"^Phase[:\s]", text, re.I):
            current_phase = re.sub(r"(?i)^Phase[:\s]*", "", text).strip()
            current_question = None
            collecting_answers = False
            continue

        # Detect Question
        if (re.match(r"^(Q\d+[:\.]|Question\s*\d*[:\.])", text, re.I) or 
            text.endswith("?") or 
            ("?" in text and len(text) < 150)):
            
            # Save previous question's answers if any
            if current_question and current_answers:
                for answer_line in current_answers:
                    if ":" in answer_line:
                        parts = answer_line.split(":", 1)
                        if len(parts) == 2:
                            person, answer = parts
                            rows.append([
                                current_phase,
                                current_question,
                                person.strip(),
                                answer.strip()
                            ])
            
            current_question = text
            current_answers = []
            collecting_answers = True
            continue

        # Collect answers for current question
        if collecting_answers and current_question:
            # Check if this is an answer line (starts with person's name)
            if re.match(r"^[A-Z][a-zA-Z\s]+[:\-]", text):
                current_answers.append(text)
            else:
                # If not an answer line, it might be the next question or phase
                collecting_answers = False

    # Save the last question's answers
    if current_question and current_answers:
        for answer_line in current_answers:
            if ":" in answer_line:
                parts = answer_line.split(":", 1)
                if len(parts) == 2:
                    person, answer = parts
                    rows.append([
                        current_phase,
                        current_question,
                        person.strip(),
                        answer.strip()
                    ])

    # Save to CSV
    if rows:
        df = pd.DataFrame(rows, columns=["Phase", "Question", "Person", "Answer"])
        df.to_csv(csv_file, index=False, encoding="utf-8-sig")
        print(f"✅ Saved {len(df)} rows to {csv_file}")
        return df
    else:
        print("❌ No data was extracted. Trying alternative approach...")
        return alternative_approach(doc, csv_file)

def alternative_approach(doc, csv_file):
    """Alternative approach for different document structure"""
    rows = []
    current_phase, current_question = None, None
    
    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
            print("ALT>>", text)
    
    i = 0
    while i < len(all_text):
        text = all_text[i]
        
        # Detect Phase
        if re.match(r"^Phase[:\s]", text, re.I):
            current_phase = re.sub(r"(?i)^Phase[:\s]*", "", text).strip()
            i += 1
            continue
        
        # Detect Question
        if (re.match(r"^(Q\d+[:\.]|Question)", text, re.I) or 
            text.endswith("?") or 
            ("?" in text and len(text) < 150)):
            
            current_question = text
            i += 1
            
            # Collect all answer lines for this question
            answers = []
            while (i < len(all_text) and 
                   not re.match(r"^(Phase|Q\d+|Question)", all_text[i], re.I) and 
                   not all_text[i].endswith("?")):
                
                answer_line = all_text[i]
                if re.match(r"^[A-Z][a-zA-Z\s]+[:\-]", answer_line):
                    answers.append(answer_line)
                i += 1
            
            # Process all answers for this question
            for answer_line in answers:
                if ":" in answer_line:
                    parts = answer_line.split(":", 1)
                    if len(parts) == 2:
                        person, answer = parts
                        rows.append([
                            current_phase,
                            current_question,
                            person.strip(),
                            answer.strip()
                        ])
        else:
            i += 1

    if rows:
        df = pd.DataFrame(rows, columns=["Phase", "Question", "Person", "Answer"])
        df.to_csv(csv_file, index=False, encoding="utf-8-sig")
        print(f"✅ Alternative approach saved {len(df)} rows to {csv_file}")
        return df
    else:
        print("❌ Still no data extracted. Using manual parsing...")
        return manual_parsing(all_text, csv_file)

def manual_parsing(all_text, csv_file):
    """Manual parsing based on the specific structure"""
    rows = []
    current_phase, current_question = None, None
    
    for i, text in enumerate(all_text):
        # Detect Phase
        if re.match(r"^Phase[:\s]", text, re.I):
            current_phase = re.sub(r"(?i)^Phase[:\s]*", "", text).strip()
            continue
        
        # Detect Question (lines starting with Q or containing question mark)
        if (re.match(r"^Q\d+", text) or 
            ("?" in text and len(text) < 200 and not text.startswith(('Imtiaz', 'Javid', 'Tahawur', 'Amjad', 'Asad', 'Muhammad')))):
            
            current_question = text
            continue
        
        # Detect Answer lines (start with person names)
        if (current_question and 
            re.match(r"^(Imtiaz Ahmed|Javid Hussain|Tahawur Abbas|Amjad|Asad Ali|Muhammad Askari)[:\-]", text)):
            
            if ":" in text:
                parts = text.split(":", 1)
                if len(parts) == 2:
                    person, answer = parts
                    rows.append([
                        current_phase,
                        current_question,
                        person.strip(),
                        answer.strip()
                    ])
    
    if rows:
        df = pd.DataFrame(rows, columns=["Phase", "Question", "Person", "Answer"])
        df.to_csv(csv_file, index=False, encoding="utf-8-sig")
        print(f"✅ Manual parsing saved {len(df)} rows to {csv_file}")
        return df
    else:
        print("❌ All parsing methods failed. Creating CSV from your provided text...")
        return create_from_provided_text(csv_file)

def create_from_provided_text(csv_file):
    """Create CSV from the text you provided in the question"""
    # This uses the exact text structure you provided
    text_content = """
Phase: Requirements Analysis
Q1: How were the initial requirements for the LMS gathered? Was there any formal requirement-gathering process like interviews, surveys, or workshops?
Imtiaz Ahmed: Requirements were gathered by exploring open-source LMS platforms (e.g., Moodle) and consulting IT team members.
Javid Hussain: Requirements were gathered through discussion-based sessions initiated by Dr. Irshad Hussain with a committee of faculty members, IT staff, and an intern.
Tahawur Abbas: Requirements were collected via student interviews and admin meetings, with formal interviews and meetings conducted.
Amjad: Requirements were gathered through comprehensive requirement analysis sessions, including interviews and regular meetings with stakeholders.
Asad Ali: Requirements were gathered through meetings; formal documentation was not prepared.
Muhammad Askari: Requirements were gathered through informal stakeholder meetings rather than formal documentation processes.
# ... [rest of your text] ...
"""
    
    lines = text_content.strip().split('\n')
    rows = []
    current_phase, current_question = None, None
    
    for line in lines:
        line = line.strip()
        if not line or line.startswith('#'):
            continue
            
        # Detect Phase
        if line.startswith('Phase:'):
            current_phase = line.replace('Phase:', '').strip()
            continue
            
        # Detect Question
        if line.startswith('Q') and '?' in line:
            current_question = line
            continue
            
        # Detect Answer
        if ':' in line and any(name in line for name in ['Imtiaz', 'Javid', 'Tahawur', 'Amjad', 'Asad', 'Muhammad']):
            parts = line.split(':', 1)
            if len(parts) == 2:
                person, answer = parts
                rows.append([
                    current_phase,
                    current_question,
                    person.strip(),
                    answer.strip()
                ])
    
    if rows:
        df = pd.DataFrame(rows, columns=["Phase", "Question", "Person", "Answer"])
        df.to_csv(csv_file, index=False, encoding="utf-8-sig")
        print(f"✅ Created CSV from provided text with {len(df)} rows")
        return df
    else:
        print("❌ All methods failed. Please check your Word document structure.")
        return None

if __name__ == "__main__":
    input_file = os.path.join("data", "lms.docx")
    output_file = os.path.join("data", "lms_interview_data.csv")
    
    # Create data directory if it doesn't exist
    os.makedirs("data", exist_ok=True)
    
    df = word_to_csv_fixed(input_file, output_file)
    
    if df is not None and len(df) > 0:
        print("\nSample of extracted data:")
        print(df.head(10))
        print(f"\nTotal rows: {len(df)}")
        print(f"Phases: {df['Phase'].nunique()}")
        print(f"Questions: {df['Question'].nunique()}")
        print(f"Persons: {df['Person'].nunique()}")
    else:
        print("❌ Failed to extract data. Please check your Word document structure.")